import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import CACHE_DIR, REPORTS_DIR


def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_metric_row(doc, label, value):
    para = doc.add_paragraph()
    run1 = para.add_run(f"{label}: ")
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = para.add_run(str(value))
    run2.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(2)
    return para


def generate_docx_report(ticker: str, financials: dict,
                          valuation: dict, news: dict,
                          memo: dict) -> Path:
    """Generate a professional DOCX investment report."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

    # ── TITLE PAGE ──────────────────────────────────────
    title = doc.add_heading(f"Investment Research Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 70, 127)
        run.font.size = Pt(24)

    subtitle = doc.add_paragraph(f"{ticker} — Apple Inc.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True

    date_para = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Recommendation badge
    rec   = valuation.get("recommendation", {})
    rec_r = rec.get("recommendation", "HOLD")
    color_map = {"BUY": (0,128,0), "HOLD": (255,140,0), "SELL": (200,0,0)}
    rec_color = color_map.get(rec_r, (128,128,128))

    rec_para = doc.add_paragraph()
    rec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rec_run = rec_para.add_run(f"RECOMMENDATION: {rec_r}")
    rec_run.bold = True
    rec_run.font.size = Pt(18)
    rec_run.font.color.rgb = RGBColor(*rec_color)

    doc.add_page_break()

    # ── SECTION 1: FINANCIAL HIGHLIGHTS ─────────────────
    add_heading(doc, "1. Financial Highlights", 1, (0, 70, 127))

    m = financials.get("metrics", {})
    r = financials.get("ratios", {})

    metrics = [
        ("Revenue",            m.get("Revenue", [{}])[0].get("formatted", "N/A")),
        ("Net Income",         m.get("NetIncome", [{}])[0].get("formatted", "N/A")),
        ("Operating Income",   m.get("OperatingIncome", [{}])[0].get("formatted", "N/A")),
        ("EPS (Diluted)",      m.get("EPS_Diluted", [{}])[0].get("formatted", "N/A")),
        ("Total Assets",       m.get("TotalAssets", [{}])[0].get("formatted", "N/A")),
        ("Cash & Equivalents", m.get("CashAndEquivalents", [{}])[0].get("formatted", "N/A")),
        ("Operating Cash Flow",m.get("OperatingCashFlow", [{}])[0].get("formatted", "N/A")),
    ]

    for label, value in metrics:
        add_metric_row(doc, label, value)

    doc.add_paragraph()
    add_heading(doc, "Key Ratios", 2)

    ratios = [
        ("Net Profit Margin",  r.get("NetProfitMargin", "N/A")),
        ("Return on Assets",   r.get("ROA", "N/A")),
        ("Return on Equity",   r.get("ROE", "N/A")),
        ("Debt / Equity",      r.get("DebtToEquity", "N/A")),
    ]
    for label, value in ratios:
        add_metric_row(doc, label, value)

    # ── SECTION 2: VALUATION ────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "2. Valuation Analysis", 1, (0, 70, 127))

    mkt = valuation.get("market", {})
    dcf = valuation.get("dcf", {})

    valuation_data = [
        ("Current Market Price",  f"${mkt.get('current_price', 0):,.2f}"),
        ("Market Capitalization", f"${mkt.get('market_cap', 0)/1e12:.2f}T"),
        ("P/E Ratio",             f"{mkt.get('pe_ratio', 0):.1f}x"),
        ("52-Week High",          f"${mkt.get('52w_high', 0):,.2f}"),
        ("52-Week Low",           f"${mkt.get('52w_low', 0):,.2f}"),
        ("DCF Intrinsic Value",   f"${dcf.get('intrinsic_value', 0):,.2f}"),
        ("DCF Upside/Downside",   f"{dcf.get('upside_downside', 0):+.1f}%"),
        ("Analyst Target Price",  f"${rec.get('analyst_target', 0):,.2f}"),
        ("Analyst Upside",        f"{rec.get('analyst_upside', 0):+.1f}%"),
    ]
    for label, value in valuation_data:
        add_metric_row(doc, label, value)

    # ── SECTION 3: NEWS SENTIMENT ────────────────────────
    doc.add_paragraph()
    add_heading(doc, "3. News Sentiment Analysis", 1, (0, 70, 127))

    sen = news.get("sentiment", {})
    breakdown = sen.get("breakdown", {})

    add_metric_row(doc, "Overall Sentiment", f"{sen.get('overall','N/A')} ({sen.get('score',0):+.3f})")
    add_metric_row(doc, "Bullish Articles",  breakdown.get("bullish", 0))
    add_metric_row(doc, "Neutral Articles",  breakdown.get("neutral", 0))
    add_metric_row(doc, "Bearish Articles",  breakdown.get("bearish", 0))

    doc.add_paragraph()
    doc.add_paragraph("Recent Headlines:").runs[0].bold = True
    for article in news.get("articles", [])[:5]:
        score = article.get("sentiment_score", 0)
        icon  = "▲" if score > 0.15 else "▼" if score < -0.15 else "–"
        p = doc.add_paragraph(
            f"{icon} [{article.get('published','')[:8]}] "
            f"{article.get('title','')[:80]}",
            style="List Bullet"
        )
        p.runs[0].font.size = Pt(10)

    # ── SECTION 4: INVESTMENT MEMO ───────────────────────
    doc.add_paragraph()
    add_heading(doc, "4. Investment Memo", 1, (0, 70, 127))

    memo_text = memo.get("memo", "")
    # Split memo into paragraphs and add each
    for line in memo_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Detect section headers (bold markdown)
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("---"):
            doc.add_paragraph()
        else:
            clean = line.replace("**", "")
            p = doc.add_paragraph(clean)
            p.runs[0].font.size = Pt(10) if p.runs else Pt(10)

    # ── FOOTER ──────────────────────────────────────────
    doc.add_page_break()
    footer_para = doc.add_paragraph(
        "This report was generated automatically by the Agentic Financial Intelligence Platform. "
        "Data sourced from SEC EDGAR, Alpha Vantage, and yfinance. "
        "This is not financial advice."
    )
    footer_para.runs[0].font.size   = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    filename = f"{ticker}_Investment_Report_{datetime.now().strftime('%Y%m%d')}.docx"
    out_path = REPORTS_DIR / filename
    doc.save(str(out_path))

    return out_path


def run(ticker: str = "AAPL") -> dict:
    """Main Report Generation Agent runner."""
    print(f"\n{'='*55}")
    print(f"  Report Generation Agent — {ticker}")
    print(f"{'='*55}\n")

    # Load all cached data
    fin_file  = CACHE_DIR / f"{ticker}_financials.json"
    val_file  = CACHE_DIR / f"{ticker}_valuation.json"
    news_file = CACHE_DIR / f"{ticker}_news.json"
    memo_file = CACHE_DIR / f"{ticker}_investment_memo.json"

    for f in [fin_file, val_file, news_file, memo_file]:
        if not f.exists():
            raise FileNotFoundError(f"Missing: {f.name} — run previous agents first")

    financials = json.loads(fin_file.read_text())
    valuation  = json.loads(val_file.read_text())
    news       = json.loads(news_file.read_text())
    memo       = json.loads(memo_file.read_text())

    print(f"  Generating DOCX report for {ticker}...")
    docx_path = generate_docx_report(ticker, financials, valuation, news, memo)

    print(f"  ✓ DOCX report saved: {docx_path.name}")
    print(f"  📁 Location: {docx_path}")

    print(f"\n{'='*55}")
    print(f"  SUCCESS — Report generated!")
    print(f"  Open: {docx_path}")
    print(f"{'='*55}\n")

    return {
        "ticker":     ticker,
        "docx_path":  str(docx_path),
        "generated":  datetime.now().isoformat()
    }


if __name__ == "__main__":
    run("AAPL")